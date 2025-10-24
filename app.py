


import tempfile
import os 
from dotenv import load_dotenv
from langdetect import detect
from telegram import Update
from telegram.ext import ApplicationBuilder, MessageHandler, CommandHandler, filters, ContextTypes
from openai import OpenAI
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# 🌿 Load environment variables
load_dotenv()

#Get API keys from environment variables
BOT_TOKEN = os.getenv("BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# 🔑 Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# 👋 Start command
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Hello! I’m your *Multilingual Meeting Assistant Bot*.\n\n"
        "🎤 Send me a voice note (English, Swahili, or Kamba) or text.\n"
        "📋 I’ll create professional English meeting minutes for you."
    )

# 🗣️ Handle audio input
async def handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("🎧 Processing your audio...")

    # Download the audio file
    file = await update.message.voice.get_file()
    file_path = await file.download_to_drive()

    # Step 1: Transcribe
    with open(file_path, "rb") as audio:
        transcript = client.audio.transcriptions.create(model="whisper-1", file=audio)
    text = transcript.text.strip()

    # Step 2: Language detection
    lang = detect(text)
    await update.message.reply_text(f"🈶 Detected language: {lang.upper()}")

    # Step 3: Translate to English if not already
    english_text = translate_to_english(text, lang)

    # Step 4: Summarize into minutes
    minutes = summarize_text(english_text)

    # Step 5: Send results
    await send_minutes(update, minutes)

# 📝 Handle text input
async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    lang = detect(text)
    await update.message.reply_text(f"📖 Detected language: {lang.upper()}")
    english_text = translate_to_english(text, lang)
    minutes = summarize_text(english_text)
    await send_minutes(update, minutes)

# 🌐 Translation helper
def translate_to_english(text, lang):
    if lang == "en":
        return text  # already English

    prompt = f"Translate the following {lang} text to English:\n\n{text}"
    response = client.chat.completions.create(
        model="gpt-5",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# 🧠 Summarize helper
def summarize_text(text):
    prompt = f"""
    Write professional and well-structured meeting minutes from the following text.
    Include:
    - Date and Time
    - Attendees (if mentioned)
    - Agenda
    - Key Points Discussed
    - Decisions and Action Items
    Make it sound natural and clear.
    Transcript:
    {text}
    """
    response = client.chat.completions.create(
        model="gpt-5",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# 📄 Generate DOCX and PDF
async def send_minutes(update, minutes_text):
    # DOCX
    doc = Document()
    doc.add_heading("Minutes of the Meeting", level=1)
    doc.add_paragraph(minutes_text)
    docx_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(docx_temp.name)

    # PDF
    pdf_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    styles = getSampleStyleSheet()
    story = [Paragraph("Minutes of the Meeting", styles["Title"]), Spacer(1, 12)]
    story.append(Paragraph(minutes_text, styles["Normal"]))
    pdf = SimpleDocTemplate(pdf_temp.name)
    pdf.build(story)

    # Send results
    await update.message.reply_text("✅ Here are your meeting minutes:")
    await update.message.reply_document(open(docx_temp.name, "rb"), filename="Meeting_Minutes.docx")
    await update.message.reply_document(open(pdf_temp.name, "rb"), filename="Meeting_Minutes.pdf")

# 🚀 Launch bot
app = ApplicationBuilder().token(BOT_TOKEN).build()
app.add_handler(CommandHandler("start", start))
app.add_handler(MessageHandler(filters.VOICE, handle_audio))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

print("🤖 Multilingual Meeting Assistant Bot is running...")
app.run_polling()

